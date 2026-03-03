from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


def set_margins(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_title(document: Document, text: str, size: int = 24) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def add_center_text(document: Document, text: str, size: int = 12, bold: bool = False) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(item, style="List Number")
        p.runs[0].font.size = Pt(11)


def add_cover_page(document: Document) -> None:
    document.add_paragraph("\n" * 6)
    add_title(document, "PRIMEX LTD", size=28)
    add_center_text(document, "Tenant Acquisition & Pilot Program Proposal", size=18, bold=True)
    add_center_text(document, "Integrated AI-Powered Enterprise Solutions", size=13)
    document.add_paragraph("\n" * 10)
    add_center_text(document, "Prepared for: Prospective Pilot Tenants", size=11)
    add_center_text(document, "Date: March 2026", size=11)
    document.add_page_break()


def add_executive_summary(document: Document) -> None:
    add_heading(document, "Executive Summary", level=1)
    add_paragraph(
        document,
        "Primex ERP is an integrated multi-tenant business platform that combines front-office selling, "
        "back-office control, and executive visibility into one connected system. The ecosystem is designed "
        "for retail, wholesale, bars, restaurants, and hybrid distribution operations that need real-time "
        "coordination across sales, stock, staff, and financial operations.",
    )
    add_paragraph(
        document,
        "At core level, Primex unifies POS transactions, inventory movements, outlet operations, permissions, "
        "and reporting in a single operational backbone. The AI-driven analytics vision strengthens this foundation "
        "with predictive insights, anomaly visibility, and actionable KPI dashboards to improve decision speed and "
        "reduce operational blind spots.",
    )
    add_paragraph(
        document,
        "The platform is positioned to directly solve high-cost operational issues including stock loss and theft, "
        "reporting gaps across branches, delayed management visibility, payroll inaccuracies, and reconciliation "
        "errors between sales and inventory. For growth-focused businesses, Primex creates a scalable operating "
        "model with tighter controls, stronger accountability, and measurable profitability gains.",
    )


def add_scoped_platform_snapshot(document: Document) -> None:
    add_heading(document, "Current Product Scope Snapshot (Frontend + Backend)", level=1)
    add_paragraph(
        document,
        "The following capability scope is derived from the current PrimePOS/PrimeERP codebase and implementation documents.",
    )

    add_heading(document, "Backend Capability Scope", level=2)
    add_bullets(
        document,
        [
            "Multi-tenant architecture with strict tenant isolation, role-based access control, and feature permissions per tenant.",
            "Core domains implemented: sales, POS modes, products, inventory, customers, shifts, outlets, staff, suppliers, reports, and notifications.",
            "Distribution/Fleet foundations: vehicle registry, driver profiles, delivery orders, trip tracking, vehicle locations, and proof-of-delivery support.",
            "Office permissions already modeled for Accounting, HR & Payroll, Reports, and Analytics access controls.",
            "Operational integrity patterns include stock movement tracking, outlet-level filtering, and audit-ready transaction history.",
        ],
    )

    add_heading(document, "Frontend Capability Scope", level=2)
    add_bullets(
        document,
        [
            "Dedicated POS experiences for retail, restaurant, bar, and single-product workflows.",
            "Dashboard modules for sales, inventory, office operations, distribution, reports, settings, and outlet-level management.",
            "Permission-gated route access aligned with tenant feature flags for app-level and feature-level controls.",
            "Advanced inventory UX including multi-unit selling, stock thresholds, import/export tools, and clearer cashier selling flows.",
            "Commercial onboarding and packaging foundations including pricing/plan guidance and onboarding flow support.",
        ],
    )


def add_segment_section(
    document: Document,
    title: str,
    pain_points: list[str],
    solutions: list[str],
    benefits: list[str],
    roi: str,
) -> None:
    add_heading(document, title, level=2)

    add_paragraph(document, "Pain Points", bold=True)
    add_bullets(document, pain_points)

    add_paragraph(document, "How Primex Solves Them", bold=True)
    add_bullets(document, solutions)

    add_paragraph(document, "Clear Benefits", bold=True)
    add_bullets(document, benefits)

    add_paragraph(document, "ROI Explanation", bold=True)
    add_paragraph(document, roi)


def add_target_segments(document: Document) -> None:
    add_heading(document, "Target Segments", level=1)

    add_segment_section(
        document,
        "Retail Stores",
        pain_points=[
            "Frequent stockouts and overstock due to weak stock visibility.",
            "Cashier errors and inconsistent checkout speed.",
            "Difficult branch-level performance tracking.",
        ],
        solutions=[
            "Real-time stock control with low-stock alerts and outlet filtering.",
            "Fast POS workflow with role controls, returns support, and clear receipts.",
            "Unified dashboards for branch and product performance.",
        ],
        benefits=[
            "Higher shelf availability and fewer lost sales.",
            "Reduced till errors and stronger cashier accountability.",
            "Faster management actions based on daily data.",
        ],
        roi=(
            "Retail tenants typically recover value through fewer stock losses, faster checkout throughput, "
            "and reduced reconciliation time. Even modest reductions in stock variance and cashier leakage can "
            "offset subscription costs rapidly and improve gross margin consistency."
        ),
    )

    add_segment_section(
        document,
        "Wholesale Businesses",
        pain_points=[
            "Bulk pricing complexity and inconsistent order processing.",
            "Poor visibility into supplier cycles and purchase planning.",
            "Manual reporting delays for management and finance teams.",
        ],
        solutions=[
            "Flexible product and unit handling for bulk sales operations.",
            "Supplier, stock movement, and transfer visibility in one platform.",
            "Consolidated reporting across sales, inventory, and operations.",
        ],
        benefits=[
            "Improved order accuracy and lower fulfillment errors.",
            "Better demand planning and reduced dead stock.",
            "Stronger controls for high-volume environments.",
        ],
        roi=(
            "Wholesale ROI comes from improved stock turnover, lower picking/dispatch errors, and faster order-to-cash cycles. "
            "As volume grows, automation and control prevent hidden margin erosion and protect working capital."
        ),
    )

    add_segment_section(
        document,
        "Bars",
        pain_points=[
            "High shrinkage risk from fast-moving drinks and poor tab controls.",
            "Inconsistent shift handovers and cash variance issues.",
            "Limited visibility into high-margin vs low-margin items.",
        ],
        solutions=[
            "Bar-specific POS workflows with rapid order and tab handling.",
            "Shift and transaction tracking with role-based authorization.",
            "Sales and item performance reporting to optimize drink mix.",
        ],
        benefits=[
            "Reduced leakage and tighter cash discipline.",
            "Faster service during peak hours.",
            "Improved profitability from smarter menu decisions.",
        ],
        roi=(
            "Bar tenants gain ROI from reduced pour-loss leakage, better shift accountability, and increased peak-hour throughput. "
            "Small improvements in control and speed typically produce significant monthly profit impact."
        ),
    )

    add_segment_section(
        document,
        "Restaurants",
        pain_points=[
            "Order delays between front-of-house and kitchen workflows.",
            "Difficulty tracking table turnover and service performance.",
            "Food cost pressure from weak inventory discipline.",
        ],
        solutions=[
            "Restaurant-focused POS workflows and order routing support.",
            "Table and order management with role-specific operations.",
            "Inventory and consumption visibility linked to sales activity.",
        ],
        benefits=[
            "Improved service speed and customer experience.",
            "Lower wastage and better portion-cost control.",
            "More predictable daily and weekly margins.",
        ],
        roi=(
            "Restaurant ROI is driven by better table productivity, reduced food waste, and improved labor efficiency. "
            "Tighter operational alignment between sales and stock directly improves contribution margin."
        ),
    )

    add_segment_section(
        document,
        "Hybrid Businesses (Retail + Bar / Wholesale + Distribution)",
        pain_points=[
            "Disconnected tools across storefront, warehouse, and delivery teams.",
            "Complex permission and control requirements by department.",
            "No unified source of truth for operational and financial decisions.",
        ],
        solutions=[
            "Single multi-module platform with tenant-specific feature activation.",
            "Hybrid support across POS, inventory, distribution/fleet, and office controls.",
            "Centralized dashboards and reporting for management visibility.",
        ],
        benefits=[
            "Cross-team coordination improves immediately.",
            "Lower software sprawl and reduced integration burden.",
            "Scalable operating model for expansion into new branches or routes.",
        ],
        roi=(
            "Hybrid operators realize ROI through system consolidation, reduced duplicate work, and fewer cross-department losses. "
            "A unified data layer improves execution quality and shortens decision cycles at management level."
        ),
    )


def add_module_overview(document: Document) -> None:
    add_heading(document, "Prime ERP Modules Overview", level=1)

    add_heading(document, "PrimePOS", level=2)
    add_bullets(
        document,
        [
            "Retail, bar, and restaurant transaction workflows.",
            "Multi-payment operations and receipt handling.",
            "Outlet-aware operations with role and permission controls.",
        ],
    )

    add_heading(document, "PrimeHR & Payroll", level=2)
    add_bullets(
        document,
        [
            "Staff and user-role operations with centralized access governance.",
            "HR/payroll feature access controlled at tenant level.",
            "Foundation for payroll accuracy through cleaner attendance/shift-linked operational records.",
        ],
    )

    add_heading(document, "PrimeAccounting", level=2)
    add_bullets(
        document,
        [
            "Office accounting controls with feature-level tenant permissions.",
            "Expense and reporting integration with operational modules.",
            "Foundation for cleaner month-end reconciliation and management finance reporting.",
        ],
    )

    add_heading(document, "PrimeFleet", level=2)
    add_bullets(
        document,
        [
            "Vehicle, driver, delivery order, and trip lifecycle management.",
            "Proof-of-delivery and vehicle location tracking readiness.",
            "Cost and profitability visibility for distribution operations.",
        ],
    )

    add_heading(document, "PrimeAnalytics", level=2)
    add_bullets(
        document,
        [
            "Management dashboards for sales, stock, outlet, and operational performance.",
            "Permission-gated analytics visibility by tenant and role.",
            "AI-driven direction for predictive insights and anomaly detection.",
        ],
    )


def add_pilot_offer(document: Document) -> None:
    add_heading(document, "Pilot Testing Offer", level=1)
    add_paragraph(
        document,
        "PRIMEX LTD proposes a structured pilot program designed to derisk adoption while proving measurable operational impact.",
    )

    add_heading(document, "Free Trial Structure", level=2)
    add_numbered(
        document,
        [
            "Week 1: Discovery and operational baseline capture.",
            "Week 2: System configuration, data import, and user role setup.",
            "Week 3: Live pilot execution with active support and KPI monitoring.",
            "Week 4: Impact review, ROI scorecard, and scale plan recommendation.",
        ],
    )

    add_heading(document, "Included in the Pilot", level=2)
    add_bullets(
        document,
        [
            "Guided onboarding support with implementation specialist involvement.",
            "Staff training for managers, cashiers, and operations teams.",
            "Custom configuration aligned to business type and outlet model.",
            "Data migration support from spreadsheets or legacy systems.",
            "Weekly checkpoint reviews with actionable improvement steps.",
        ],
    )


def add_sales_scripts(document: Document) -> None:
    add_heading(document, "Sales Scripts", level=1)

    add_heading(document, "Cold Outreach Script", level=2)
    add_paragraph(
        document,
        "Hello [Name], this is [Your Name] from PRIMEX LTD. We help retail and distribution businesses reduce stock loss, "
        "tighten cash control, and improve branch visibility using one integrated platform. Most teams we speak with are "
        "using separate tools for POS, inventory, and reporting, which creates blind spots and avoidable losses. "
        "Would you be open to a 20-minute session this week to see how a pilot could improve your current operations?"
    )

    add_heading(document, "WhatsApp Outreach Message", level=2)
    add_paragraph(
        document,
        "Hi [Name], I hope you are well. I’m reaching out from PRIMEX LTD. We run an integrated POS + inventory + reporting "
        "platform built for businesses like yours. We are currently onboarding pilot tenants and can help you reduce stock leakage, "
        "improve daily reporting, and strengthen staff accountability. If useful, I can share a short demo slot this week."
    )

    add_heading(document, "Follow-Up Script", level=2)
    add_paragraph(
        document,
        "Hi [Name], just following up on my previous message. We’re helping teams replace disconnected systems with one integrated "
        "operating platform and measurable KPI improvements in the first month. If timing is better now, I can reserve a demo and "
        "prepare a pilot plan specific to your store/branch structure."
    )

    add_heading(document, "Demo Closing Script", level=2)
    add_paragraph(
        document,
        "Based on your current setup, the fastest and lowest-risk next step is a structured pilot. We will configure the system around "
        "your workflows, train your team, migrate your key data, and track agreed KPIs. At the end of the pilot, you get a clear "
        "business case with ROI evidence before any long-term commitment. If you approve, we can schedule pilot kickoff immediately."
    )


def add_objection_handling(document: Document) -> None:
    add_heading(document, "Objection Handling", level=1)

    add_heading(document, "Objection: “We already use another POS”", level=2)
    add_paragraph(
        document,
        "That makes sense, and we are not asking you to take blind risk. Most of our strongest tenants came from existing tools that "
        "handled checkout but lacked integrated control across inventory, reporting, and operations. Our pilot lets you validate measurable "
        "improvements before making a full transition decision."
    )

    add_heading(document, "Objection: “It’s too expensive”", level=2)
    add_paragraph(
        document,
        "The real comparison is not subscription vs no subscription, but controlled operations vs silent losses. Stock leakage, "
        "cash variance, reporting delays, and avoidable payroll or reconciliation errors usually cost more than software. "
        "We structure pilots to prove value first, then align package cost to delivered impact."
    )

    add_heading(document, "Objection: “We don’t trust new systems”", level=2)
    add_paragraph(
        document,
        "That concern is valid. We reduce adoption risk with phased onboarding, team training, controlled rollout, and close support. "
        "You keep visibility throughout the pilot, and decisions are based on performance data, not assumptions."
    )

    add_heading(document, "Objection: “We are not ready”", level=2)
    add_paragraph(
        document,
        "Readiness is exactly why we run pilots. We start from your current maturity level, migrate only the right data first, "
        "and help your team build confidence through guided execution. You don’t need perfect readiness to start; you need a structured plan."
    )


def add_closing(document: Document) -> None:
    add_heading(document, "Strategic Close", level=1)
    add_paragraph(
        document,
        "PRIMEX LTD is positioned to become the operational system of record for growth-focused businesses that need control, speed, "
        "and scale. This proposal is designed to help qualified tenants adopt with low risk, measurable gains, and executive-level confidence.",
    )
    add_paragraph(
        document,
        "Next Step: Approve pilot discovery meeting and finalize pilot scope, timeline, and KPI success metrics.",
        bold=True,
    )


def generate_document(output_filename: str = "Primex_Tenant_Acquisition_Marketing_Pitch.docx") -> None:
    document = Document()
    set_margins(document)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    add_cover_page(document)
    add_executive_summary(document)
    add_scoped_platform_snapshot(document)
    add_target_segments(document)
    add_module_overview(document)
    add_pilot_offer(document)
    add_sales_scripts(document)
    add_objection_handling(document)
    add_closing(document)

    document.save(output_filename)
    print(f"Document created successfully: {output_filename}")


if __name__ == "__main__":
    generate_document()